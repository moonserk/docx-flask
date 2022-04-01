import os

from flask import Flask, url_for, render_template, request, redirect, session, send_file
from werkzeug.utils import secure_filename

import pyexcel as pe
from docx import Document

from doc import create_documents

UPLOAD_FOLDER = './uploads'
ALLOWED_EXTENSIONS = {'xls', 'docx'}

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = '73870e7f-634d-433b-946a-8d20132bafac'

@app.route('/')
def home():
    url_for('static', filename="style.css")
    return render_template('template.html')

@app.route('/documents', methods=['GET', 'POST'])
def documents():
    items = []
    if (request.method == 'POST'):
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)

        file = request.files['file']

        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file.save(os.path.join('./uploads', filename))

        return redirect(url_for('process', filename=filename))

    return render_template('documents.html')

@app.route('/documents/<filename>', methods=['GET', 'POST'])
def process(filename):
    sheet = pe.get_sheet(file_name='./uploads/' + filename, name_columns_by_row=0).dict
    keys = list(sheet.keys())

    if (request.method == 'POST'):
        file = request.files['file']
        document = Document(file)
        create_documents(sheet,document)

        return redirect(url_for('dir_viewer' ,path='/docs'))

    return render_template('process.html', items=sheet, keys=keys)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def combine_word_documents(files):
    merged_document = Document()

    for index, file in enumerate(files):
        sub_doc = Document(file)

        if index < len(files) - 1:
            sub_doc.add_page_break()

        for element in sub_doc.element.body:
            merged_document.element.body.append(element)

    merged_document.save('./docs/all.docx')

@app.route('/<path:path>', methods=['GET', 'POST'])
def dir_viewer(path):
    entries = os.scandir(path)
    files = []

    for filedocx in os.listdir(path):
        if filedocx.endswith(".docx"):
            files.append("./docs/" + filedocx)

    if (request.method == 'POST'):
        combine_word_documents(files)

    return render_template('files.html', entries=entries)

@app.route('/download/<path:path>')
def download_file(path):
    return send_file(path)

if __name__ == '__main__':
    app.run(environs, start_response)
